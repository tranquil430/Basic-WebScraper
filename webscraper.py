import requests
from bs4 import BeautifulSoup 
from docx import Document 


base_url = "https://novelbin.phieuvu.com/book/the-legendary-mechanic/chapter-"

# Create a new Document object for the Word file
doc = Document()

def get_paragraph(url, chapter_number):
    # Send a request to the new URL
    response = requests.get(url)
    
    # Check if request was successful
    if response.status_code == 200:
        # Parse HTML content
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Find all paragraphs
        paragraphs = soup.find_all('p')
        
        # Add chapter title to the document
        doc.add_heading(f'Chapter {chapter_number}', level=1)
        
        # Add each paragraph to the document
        for idx, paragraph in enumerate(paragraphs, 1):
            doc.add_paragraph(paragraph.get_text())
    else:
        print(f"Failed to retrieve the webpage. Status code: {response.status_code}")

# Number of chapters in the novel
chapter_count = 5
x = 1

while x <= chapter_count:

    new_url = base_url + str(x)
    get_paragraph(new_url, x)
    x += 1

doc.save('PATH')

